"""
GovKnowledge AI — The Intelligent Knowledge Hub for Government Officers
Streamlit + LlamaIndex + Ollama (free) or OpenAI
"""

from __future__ import annotations

import base64
import hashlib
import logging
import os
import re
import shutil
import time
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Any

import streamlit as st
from dotenv import load_dotenv

load_dotenv()

# ---------------------------------------------------------------------------
# Logging (console only — never shown to end users)
# ---------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Paths & constants
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent
DATA_SOURCE_DIR = BASE_DIR / "data_source"
INDEX_CACHE_DIR = BASE_DIR / ".index_cache"
PERSIST_DIR = INDEX_CACHE_DIR / "storage"
FINGERPRINT_FILE = INDEX_CACHE_DIR / "corpus.fp"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

GROQ_VISION_MODEL_DEFAULT = "qwen/qwen3.6-27b"
GROQ_VISION_MODEL_FALLBACKS = (
    "qwen/qwen3.6-27b",
    "meta-llama/llama-4-maverick-17b-128e-instruct",
    "meta-llama/llama-4-scout-17b-16e-instruct",
)

APP_NAME = "GovKnowledge AI"
APP_TAGLINE = "The Intelligent Knowledge Hub for Government Officers"
DEVELOPER_SITE_URL = "https://tharshan.lk"

SYSTEM_PROMPT = """You are an elite legal and regulatory compliance assistant for public sector officers.

When answering queries based on the provided documents, you must:
1. Provide accurate, concise answers grounded ONLY in the supplied context.
2. Use formal, professional language suitable for government correspondence.
3. Strictly extract and format your source citations at the TOP of your response inside a clearly marked block.

Citation format (one line per source used):
🏛️ Organization: [Name] | 📅 Year: [YYYY] | 📜 Reference Section: [Section/Paragraph]

If Organization, Year, or Section/Paragraph cannot be determined from the text, write
"Not explicitly mentioned in text" for that field.

After the citation block, provide your detailed answer. If the context does not contain
enough information, state that clearly and do not invent citations or legal references."""

GOVERNMENT_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Source+Serif+4:wght@400;600;700&family=Source+Sans+3:wght@400;500;600;700&display=swap');

:root {
    --gov-navy: #1A2B4A;
    --gov-navy-light: #2C4A7C;
    --gov-gold: #C9A227;
    --gov-gold-light: #E8C547;
    --gov-cream: #F8F9FB;
    --gov-border: #D4DCE8;
}

html, body, [class*="css"] {
    font-family: 'Source Sans 3', sans-serif;
}

.main .block-container {
    padding-top: 0 !important;
    padding-bottom: 3rem;
    max-width: 920px;
}

/* Tighten gap between top menu bar and main content */
header[data-testid="stHeader"] {
    height: auto !important;
    min-height: 0 !important;
    padding: 0.15rem 0 0.1rem 0 !important;
    margin-bottom: 0 !important;
}
[data-testid="stToolbar"] {
    top: 0.25rem !important;
}
[data-testid="stAppViewContainer"] > .main {
    padding-top: 0 !important;
}
[data-testid="stMainBlockContainer"] {
    padding-top: 0.15rem !important;
}

/* Hero header */
.gov-header {
    position: relative;
    background: linear-gradient(135deg, var(--gov-navy) 0%, var(--gov-navy-light) 100%);
    border-left: 6px solid var(--gov-gold);
    border-radius: 0 12px 12px 0;
    padding: 0.9rem 2rem 1rem 2rem;
    margin-top: 0;
    margin-bottom: 1rem;
    box-shadow: 0 4px 20px rgba(26, 43, 74, 0.15);
}
.gov-header h1 {
    font-family: 'Source Serif 4', serif;
    color: #FFFFFF !important;
    font-size: 1.85rem;
    font-weight: 700;
    margin: 0 0 0.35rem 0;
    letter-spacing: 0.02em;
    padding-right: 8.5rem;
}
.gov-header p {
    color: rgba(255, 255, 255, 0.88);
    margin: 0;
    font-size: 1rem;
    padding-right: 8.5rem;
}
.gov-site-fab {
    position: absolute;
    top: 1.35rem;
    right: 1.25rem;
    z-index: 1;
    display: inline-flex;
    align-items: center;
    gap: 0.4rem;
    padding: 0.42rem 0.9rem;
    background: rgba(255, 255, 255, 0.12);
    color: #FFFFFF !important;
    text-decoration: none !important;
    border-radius: 999px;
    font-size: 0.8rem;
    font-weight: 700;
    border: 1.5px solid var(--gov-gold);
    box-shadow: 0 2px 10px rgba(0, 0, 0, 0.15);
    letter-spacing: 0.02em;
    transition: transform 0.2s ease, background 0.2s ease;
    white-space: nowrap;
}
.gov-site-fab:hover {
    transform: translateY(-1px);
    background: rgba(255, 255, 255, 0.2);
    color: var(--gov-gold-light) !important;
}
.gov-site-fab-icon {
    font-size: 1rem;
    line-height: 1;
}

/* Status badges */
.status-badge {
    display: inline-flex;
    align-items: center;
    gap: 0.4rem;
    padding: 0.35rem 0.75rem;
    border-radius: 999px;
    font-size: 0.82rem;
    font-weight: 600;
    margin-bottom: 0.5rem;
}
.status-ok {
    background: #E6F4EA;
    color: #1B5E20;
    border: 1px solid #A5D6A7;
}
.status-warn {
    background: #FFF8E1;
    color: #E65100;
    border: 1px solid #FFE082;
}
.status-error {
    background: #FFEBEE;
    color: #B71C1C;
    border: 1px solid #EF9A9A;
}

/* Citation callout */
.citation-box {
    background: linear-gradient(to right, #FFFDF5, #FFFBEB);
    border: 1px solid var(--gov-gold);
    border-left: 5px solid var(--gov-gold);
    border-radius: 8px;
    padding: 1rem 1.25rem;
    margin: 0.75rem 0 1rem 0;
    font-size: 0.92rem;
    line-height: 1.6;
    color: var(--gov-navy);
}
.citation-box strong {
    color: var(--gov-navy);
}

/* Source document panel */
.source-docs-header {
    font-family: 'Source Serif 4', serif;
    color: var(--gov-navy);
    font-size: 1rem;
    font-weight: 600;
    margin: 1rem 0 0.5rem 0;
}
.source-doc-meta {
    color: #4A5568;
    font-size: 0.88rem;
    margin-bottom: 0.75rem;
}

.query-panel {
    background: #F8F9FB;
    border: 1px solid var(--gov-border);
    border-radius: 12px;
    padding: 1rem 1.25rem 0.5rem 1.25rem;
    margin-bottom: 1.25rem;
}
.query-panel label {
    font-weight: 600;
    color: var(--gov-navy);
}
.query-row [data-testid="column"] {
    display: flex;
    align-items: flex-end;
}

/* Chat bubbles */
.stChatMessage {
    border-radius: 12px !important;
    border: 1px solid var(--gov-border) !important;
}

/* Hide native Streamlit sidebar toggle — custom >> / << button used instead */
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapsedControl"],
[data-testid="stSidebarCollapseButton"],
[data-testid="collapsedControl"] button,
[data-testid="stSidebarCollapsedControl"] button {
    display: none !important;
    visibility: hidden !important;
    opacity: 0 !important;
    pointer-events: none !important;
}

/* Custom sidebar show/hide (controlled by gov-ui.js) */
html.gov-sidebar-hidden section[data-testid="stSidebar"] {
    transform: translateX(-110%) !important;
    min-width: 0 !important;
    max-width: 0 !important;
    width: 0 !important;
    overflow: hidden !important;
    opacity: 0 !important;
    visibility: hidden !important;
    pointer-events: none !important;
}
html.gov-sidebar-force-show section[data-testid="stSidebar"] {
    display: block !important;
    transform: none !important;
    min-width: 21rem !important;
    width: 21rem !important;
    max-width: 21rem !important;
    opacity: 1 !important;
    visibility: visible !important;
    pointer-events: auto !important;
}

/* Sidebar styling */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #F0F4FA 0%, #E8EDF5 100%);
    border-right: 1px solid var(--gov-border);
}
section[data-testid="stSidebar"] .block-container {
    padding-top: 0.65rem;
}

.sidebar-title {
    font-family: 'Source Serif 4', serif;
    color: var(--gov-navy);
    font-size: 1.1rem;
    font-weight: 700;
    border-bottom: 2px solid var(--gov-gold);
    padding-bottom: 0.5rem;
    margin-bottom: 1rem;
}

/* Custom developer footer */
.gov-footer {
    margin-top: 3rem;
    padding: 2rem 1.5rem 1.5rem;
    background: linear-gradient(180deg, #F8F9FB 0%, #EEF2F7 100%);
    border-top: 3px solid var(--gov-gold);
    border-radius: 16px 16px 0 0;
    box-shadow: 0 -4px 24px rgba(26, 43, 74, 0.08);
    text-align: center;
}
.gov-footer-inner {
    max-width: 640px;
    margin: 0 auto;
}
.gov-footer-label {
    display: block;
    font-size: 0.78rem;
    font-weight: 600;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: #6B7A90;
    margin-bottom: 0.35rem;
}
.gov-footer-name {
    display: block;
    font-family: 'Source Serif 4', serif;
    font-size: 1.45rem;
    font-weight: 700;
    color: var(--gov-navy);
    letter-spacing: 0.04em;
    margin-bottom: 0.65rem;
}
.gov-footer-tagline {
    font-family: 'Source Serif 4', serif;
    font-style: italic;
    font-size: 1.02rem;
    color: var(--gov-navy-light);
    margin: 0 0 1rem 0;
    line-height: 1.5;
}
.gov-footer-email {
    display: inline-flex;
    align-items: center;
    gap: 0.45rem;
    padding: 0.55rem 1.15rem;
    background: linear-gradient(135deg, var(--gov-navy) 0%, var(--gov-navy-light) 100%);
    color: #FFFFFF !important;
    text-decoration: none !important;
    border-radius: 999px;
    font-size: 0.9rem;
    font-weight: 600;
    transition: transform 0.2s ease, box-shadow 0.2s ease;
    box-shadow: 0 4px 14px rgba(26, 43, 74, 0.22);
}
.gov-footer-email:hover {
    transform: translateY(-1px);
    box-shadow: 0 6px 18px rgba(26, 43, 74, 0.28);
    color: var(--gov-gold-light) !important;
}
.gov-footer-divider {
    width: 60px;
    height: 3px;
    background: linear-gradient(90deg, transparent, var(--gov-gold), transparent);
    margin: 1.25rem auto 0.85rem;
    border-radius: 2px;
}
.gov-footer-copy {
    font-size: 0.78rem;
    color: #8A96A8;
    margin: 0;
}

/* Hide Streamlit branding footer on Community Cloud */
footer { visibility: hidden; height: 0 !important; }

/* Hide Streamlit menu, profile, GitHub, Streamlit links */
[data-testid="stToolbar"],
[data-testid="stToolbarActions"],
[data-testid="stHeaderActionElements"],
#MainMenu,
[class*="viewerBadge"],
[class*="ViewerBadge"],
#GithubIcon,
a[href*="github.com"],
a[href*="github.dev"],
a[href*="streamlit.io"],
a[href*="streamlit.com"],
a[href*="share.streamlit"],
[class*="profile"],
[class*="Profile"] {
    display: none !important;
    visibility: hidden !important;
    pointer-events: none !important;
    width: 0 !important;
    height: 0 !important;
    opacity: 0 !important;
    position: absolute !important;
    left: -9999px !important;
}
</style>
"""


# ---------------------------------------------------------------------------
# Configuration helpers
# ---------------------------------------------------------------------------
def _get_secret(key: str, default: str = "") -> str:
    """Read from Streamlit secrets first, then environment variables."""
    try:
        return st.secrets[key]
    except (KeyError, FileNotFoundError, AttributeError, TypeError):
        return os.getenv(key, default)


def get_config() -> dict[str, str]:
    provider = _get_secret("AI_PROVIDER", "groq").lower().strip()
    return {
        "provider": provider,
        "openai_api_key": _get_secret("OPENAI_API_KEY"),
        "openai_model": _get_secret("OPENAI_MODEL", "gpt-4o-mini"),
        "embed_model": _get_secret("EMBED_MODEL", "text-embedding-3-small"),
        "groq_api_key": _get_secret("GROQ_API_KEY"),
        "groq_model": _get_secret("GROQ_MODEL", "llama-3.1-8b-instant"),
        "groq_vision_model": _get_secret("GROQ_VISION_MODEL", GROQ_VISION_MODEL_DEFAULT),
        "fastembed_model": _get_secret("FASTEMBED_MODEL", "BAAI/bge-small-en-v1.5"),
        "ollama_model": _get_secret("OLLAMA_MODEL", "llama3.2"),
        "ollama_embed_model": _get_secret("OLLAMA_EMBED_MODEL", "nomic-embed-text"),
        "ollama_base_url": _get_secret("OLLAMA_BASE_URL", "http://localhost:11434"),
        "admin_password": _get_secret("ADMIN_PASSWORD", "admin123"),
    }


def is_ai_ready(config: dict[str, str]) -> bool:
    if config["provider"] == "openai":
        return bool(config["openai_api_key"])
    if config["provider"] == "groq":
        return bool(config["groq_api_key"])
    return ollama_is_running(config["ollama_base_url"])


def ollama_is_running(base_url: str) -> bool:
    try:
        with urllib.request.urlopen(f"{base_url.rstrip('/')}/api/tags", timeout=3) as resp:
            return resp.status == 200
    except (urllib.error.URLError, TimeoutError, OSError):
        return False


def configure_llama_settings(config: dict[str, str]) -> None:
    from llama_index.core import Settings
    from llama_index.core.node_parser import SentenceSplitter

    Settings.node_parser = SentenceSplitter(chunk_size=2048, chunk_overlap=100)

    if config["provider"] == "openai":
        from llama_index.embeddings.openai import OpenAIEmbedding
        from llama_index.llms.openai import OpenAI

        Settings.llm = OpenAI(model=config["openai_model"], api_key=config["openai_api_key"])
        Settings.embed_model = OpenAIEmbedding(
            model=config["embed_model"],
            api_key=config["openai_api_key"],
            embed_batch_size=8,
        )
    elif config["provider"] == "groq":
        from llama_index.llms.openai import OpenAI

        Settings.embed_model = _build_groq_embed_model(config)
        Settings.llm = OpenAI(
            model=config["groq_model"],
            api_key=config["groq_api_key"],
            api_base="https://api.groq.com/openai/v1",
            is_chat_model=True,
            temperature=0.1,
            max_tokens=2048,
            context_window=8192,
        )
        Settings.node_parser = SentenceSplitter(chunk_size=1024, chunk_overlap=80)
    else:
        from llama_index.embeddings.ollama import OllamaEmbedding
        from llama_index.llms.ollama import Ollama

        Settings.llm = Ollama(
            model=config["ollama_model"],
            base_url=config["ollama_base_url"],
            request_timeout=300.0,
        )
        Settings.embed_model = OllamaEmbedding(
            model_name=config["ollama_embed_model"],
            base_url=config["ollama_base_url"],
        )


def model_cache_key(config: dict[str, str]) -> str:
    if config["provider"] == "openai":
        return f"openai:{config['openai_model']}:{config['embed_model']}"
    if config["provider"] == "groq":
        return f"groq:{config['groq_model']}:{config['fastembed_model']}"
    return f"ollama:{config['ollama_model']}:{config['ollama_embed_model']}"


# ---------------------------------------------------------------------------
# Metadata extraction from filenames & content
# ---------------------------------------------------------------------------
YEAR_PATTERN = re.compile(r"\b(19|20)\d{2}\b")
SECTION_PATTERN = re.compile(
    r"(?:section|sec\.?|paragraph|para\.?|clause|article|reg\.?)\s*[\d\.ivxlcdm\-]+",
    re.IGNORECASE,
)
ORG_KEYWORDS = [
    "ministry", "department", "commission", "authority", "council",
    "bureau", "office", "agency", "government", "public service",
    "treasury", "finance", "health", "education", "defence", "defense",
]


def _guess_organization(filename: str, text_sample: str = "") -> str:
    combined = f"{filename} {text_sample[:2000]}".lower()
    for keyword in ORG_KEYWORDS:
        match = re.search(
            rf"([\w\s\-&']+{keyword}[\w\s\-&']*)",
            combined,
            re.IGNORECASE,
        )
        if match:
            org = match.group(1).strip(" -_")
            return org.title()[:120]
    stem = Path(filename).stem.replace("_", " ").replace("-", " ")
    return stem.title()[:120] if stem else "Government Authority"


def _guess_year(filename: str, text_sample: str = "") -> str:
    for source in (filename, text_sample[:1500]):
        match = YEAR_PATTERN.search(source)
        if match:
            return match.group(0)
    return "Not explicitly mentioned in text"


def _extract_sections(text: str) -> list[str]:
    return list(dict.fromkeys(m.group(0) for m in SECTION_PATTERN.finditer(text[:8000])))[:5]


def build_document_metadata(filename: str, text: str) -> dict[str, str]:
    sections = _extract_sections(text)
    return {
        "filename": filename,
        "organization": _guess_organization(filename, text),
        "year": _guess_year(filename, text),
        "reference_section": sections[0] if sections else "Not explicitly mentioned in text",
        "source_path": str(DATA_SOURCE_DIR / filename),
    }


# ---------------------------------------------------------------------------
# Document loading
# ---------------------------------------------------------------------------
def scan_data_source() -> list[Path]:
    if not DATA_SOURCE_DIR.exists():
        DATA_SOURCE_DIR.mkdir(parents=True, exist_ok=True)
        return []
    files: list[Path] = []
    for path in sorted(DATA_SOURCE_DIR.iterdir()):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)
    return files


def _extract_pdf_with_pypdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _read_pdf_sidecar(path: Path) -> str:
    """Optional pre-extracted text: `document.pdf` → `document.ocr.txt`."""
    sidecar = path.with_name(f"{path.stem}.ocr.txt")
    if sidecar.is_file():
        return sidecar.read_text(encoding="utf-8", errors="replace")
    return ""


def _local_ocr_available() -> bool:
    try:
        import fitz  # noqa: F401
        from rapidocr_onnxruntime import RapidOCR  # noqa: F401
        return True
    except ImportError:
        return False


def _is_model_not_found_error(exc: Exception) -> bool:
    msg = str(exc).lower()
    return "model_not_found" in msg or "does not exist" in msg or "404" in msg


def _groq_vision_model_candidates(config: dict[str, str]) -> list[str]:
    preferred = (config.get("groq_vision_model") or GROQ_VISION_MODEL_DEFAULT).strip()
    candidates = [preferred]
    for model in GROQ_VISION_MODEL_FALLBACKS:
        if model not in candidates:
            candidates.append(model)
    return candidates


def _groq_vision_page_text(client: Any, model: str, encoded_png: str) -> str:
    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract ALL text from this official government document image. "
                            "Return only the extracted text. Preserve headings and numbering."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{encoded_png}"},
                    },
                ],
            }
        ],
        temperature=0.0,
        max_tokens=4096,
    )
    return (response.choices[0].message.content or "").strip()


def _extract_pdf_with_groq_vision(path: Path) -> str:
    """Cloud-friendly OCR using Groq Vision (works on Streamlit Cloud / Render)."""
    import base64

    import fitz
    from openai import OpenAI

    config = get_config()
    if not config.get("groq_api_key"):
        return ""

    client = OpenAI(
        api_key=config["groq_api_key"],
        base_url="https://api.groq.com/openai/v1",
    )
    doc = fitz.open(str(path))
    parts: list[str] = []
    active_model: str | None = None
    model_errors: list[str] = []
    try:
        for page in doc:
            png_bytes = page.get_pixmap(dpi=150).tobytes("png")
            encoded = base64.b64encode(png_bytes).decode("ascii")
            page_text = ""
            models_to_try = [active_model] if active_model else _groq_vision_model_candidates(config)
            for model in models_to_try:
                if not model:
                    continue
                try:
                    page_text = _groq_vision_page_text(client, model, encoded)
                    if active_model != model:
                        logger.info("Groq Vision OCR using model %s for %s", model, path.name)
                    active_model = model
                    break
                except Exception as exc:
                    if _is_model_not_found_error(exc):
                        model_errors.append(f"{model}: unavailable")
                        continue
                    raise
            if page_text:
                parts.append(page_text)
    finally:
        doc.close()

    if not parts and model_errors:
        logger.warning(
            "Groq Vision OCR failed for %s. Tried: %s",
            path.name,
            "; ".join(dict.fromkeys(model_errors)),
        )
    return "\n\n".join(parts)


def _extract_pdf_with_ocr(path: Path) -> str:
    """OCR fallback for scanned/image-only PDFs (no embedded text layer)."""
    import fitz
    import numpy as np
    from rapidocr_onnxruntime import RapidOCR

    ocr = RapidOCR()
    doc = fitz.open(str(path))
    parts: list[str] = []
    try:
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            if pix.n == 4:
                img = img[:, :, :3]
            result, _ = ocr(img)
            if result:
                parts.append("\n".join(line[1] for line in result))
    finally:
        doc.close()
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> tuple[str, str | None]:
    try:
        text = _extract_pdf_with_pypdf(path)
        if text.strip():
            return text, None

        sidecar_text = _read_pdf_sidecar(path)
        if sidecar_text.strip():
            logger.info("Using sidecar OCR text for %s", path.name)
            return sidecar_text, None

        logger.info("No text layer in %s — trying OCR.", path.name)

        config = get_config()
        if config.get("groq_api_key"):
            groq_text = _extract_pdf_with_groq_vision(path)
            if groq_text.strip():
                logger.info("Groq Vision OCR succeeded for %s (%d chars).", path.name, len(groq_text))
                return groq_text, None

        if _local_ocr_available():
            ocr_text = _extract_pdf_with_ocr(path)
            if ocr_text.strip():
                logger.info("Local OCR succeeded for %s (%d chars).", path.name, len(ocr_text))
                return ocr_text, None

        return "", (
            "No extractable text (scanned/image PDF). "
            "On Streamlit Cloud, ensure GROQ_API_KEY is set for vision OCR, "
            "or add a sidecar file named "
            f"{path.stem}.ocr.txt"
        )
    except Exception as exc:
        logger.exception("Failed to read PDF: %s", path.name)
        return "", str(exc)


def _read_docx(path: Path) -> tuple[str, str | None]:
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return "", "Document contains no readable paragraphs."
        return "\n\n".join(paragraphs), None
    except Exception as exc:
        logger.exception("Failed to read DOCX: %s", path.name)
        return "", str(exc)


def _read_plaintext(path: Path) -> tuple[str, str | None]:
    try:
        return path.read_text(encoding="utf-8", errors="replace"), None
    except Exception as exc:
        logger.exception("Failed to read text file: %s", path.name)
        return "", str(exc)


def load_single_document(path: Path) -> tuple[dict[str, Any] | None, str | None]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, err = _read_pdf(path)
    elif suffix in {".docx", ".doc"}:
        if suffix == ".doc":
            return None, "Legacy .doc format is not supported; please convert to .docx or .pdf."
        text, err = _read_docx(path)
    else:
        text, err = _read_plaintext(path)

    if err:
        return None, err
    if not text.strip():
        return None, "File is empty or unreadable."

    metadata = build_document_metadata(path.name, text)
    return {"text": text, "metadata": metadata}, None


def compute_corpus_fingerprint(files: list[Path]) -> str:
    hasher = hashlib.sha256()
    for path in files:
        stat = path.stat()
        hasher.update(f"{path.name}:{stat.st_size}:{stat.st_mtime}".encode())
    return hasher.hexdigest()


def documents_changed(files: list[Path]) -> bool:
    """True when data_source/ differs from the last indexed corpus."""
    if not files:
        return False
    if not st.session_state.corpus_fingerprint:
        return True
    return compute_corpus_fingerprint(files) != st.session_state.corpus_fingerprint


# ---------------------------------------------------------------------------
# LlamaIndex RAG engine (in-memory vector store — no C++ build tools needed)
# ---------------------------------------------------------------------------
def load_all_documents(files: list[Path]) -> tuple[list, list[str]]:
    from llama_index.core import Document

    documents: list[Document] = []
    errors: list[str] = []

    for path in files:
        payload, err = load_single_document(path)
        if err:
            errors.append(f"{path.name}: {err}")
            continue
        meta = payload["metadata"]
        documents.append(
            Document(
                text=payload["text"],
                metadata={
                    "filename": meta["filename"],
                    "organization": meta["organization"],
                    "year": meta["year"],
                    "reference_section": meta["reference_section"],
                },
            )
        )
    return documents, errors


MAX_CHARS_PER_DOC = 350_000  # Streamlit Cloud memory limit safeguard


def chunk_text_simple(text: str, chunk_size: int = 1200, overlap: int = 150) -> list[str]:
    text = text.strip()
    if not text:
        return []
    chunks: list[str] = []
    step = max(chunk_size - overlap, 1)
    for start in range(0, len(text), step):
        piece = text[start : start + chunk_size].strip()
        if piece:
            chunks.append(piece)
    return chunks


class GroqTfidfChatEngine:
    """Lightweight Groq RAG for Streamlit Cloud — no heavy embedding models."""

    def __init__(self, config: dict[str, str], chunks: list[str], metadatas: list[dict[str, str]]):
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity

        self.config = config
        self.chunks = chunks
        self.metadatas = metadatas
        self._cosine_similarity = cosine_similarity
        self._vectorizer = TfidfVectorizer(
            max_features=8000,
            stop_words="english",
            token_pattern=r"(?u)\b[\w\-]{2,}\b",
        )
        try:
            self._matrix = self._vectorizer.fit_transform(chunks)
        except ValueError as exc:
            if "empty vocabulary" not in str(exc).lower():
                raise
            self._vectorizer = TfidfVectorizer(max_features=8000, token_pattern=r"(?u)\b[\w\-]{2,}\b")
            self._matrix = self._vectorizer.fit_transform(chunks)
        from openai import OpenAI

        self._client = OpenAI(
            api_key=config["groq_api_key"],
            base_url="https://api.groq.com/openai/v1",
        )

    def _retrieve(self, prompt: str, top_k: int = 5) -> list[tuple[str, dict[str, str]]]:
        query_vec = self._vectorizer.transform([prompt])
        scores = self._cosine_similarity(query_vec, self._matrix).flatten()
        ranked = scores.argsort()[::-1][:top_k]
        results: list[tuple[str, dict[str, str]]] = []
        for idx in ranked:
            if scores[idx] <= 0:
                continue
            results.append((self.chunks[idx], self.metadatas[idx]))
        return results

    def chat(
        self,
        prompt: str,
        history: list[dict[str, str]] | None = None,
    ) -> tuple[str, list[dict[str, str]]]:
        retrieved = self._retrieve(prompt)
        sources = unique_source_records(meta for _, meta in retrieved)
        if not retrieved:
            context = "No matching document sections were found."
            citation_hint = ""
        else:
            context_parts = []
            citation_lines = []
            seen_citations: set[str] = set()
            for chunk, meta in retrieved:
                context_parts.append(chunk)
                cite = (
                    f"🏛️ Organization: {meta.get('organization', 'Not explicitly mentioned in text')} | "
                    f"📅 Year: {meta.get('year', 'Not explicitly mentioned in text')} | "
                    f"📜 Reference Section: {meta.get('reference_section', 'Not explicitly mentioned in text')} | "
                    f"📁 Document: {meta.get('filename', 'Unknown')}"
                )
                if cite not in seen_citations:
                    seen_citations.add(cite)
                    citation_lines.append(cite)
            context = "\n\n---\n\n".join(context_parts)
            citation_hint = "\n".join(citation_lines)

        user_prompt = (
            f"Use ONLY the context below from official documents.\n\n"
            f"CONTEXT:\n{context}\n\n"
            f"OFFICER QUESTION:\n{prompt}\n\n"
            f"Suggested source lines (verify against context):\n{citation_hint}"
        )

        messages: list[dict[str, str]] = [{"role": "system", "content": SYSTEM_PROMPT}]
        for msg in (history or [])[-6:]:
            messages.append({"role": msg["role"], "content": msg["content"]})
        messages.append({"role": "user", "content": user_prompt})

        response = self._client.chat.completions.create(
            model=self.config["groq_model"],
            messages=messages,
            temperature=0.1,
            max_tokens=2048,
        )
        content = response.choices[0].message.content
        if not content:
            raise ValueError("Groq returned an empty response.")
        return content.strip(), sources


@st.cache_resource(show_spinner=False)
def get_groq_tfidf_engine(corpus_fingerprint: str, config_snapshot: tuple):
    config = dict(config_snapshot)
    files = scan_data_source()
    if not files:
        return None, [], []

    chunks: list[str] = []
    metadatas: list[dict[str, str]] = []
    errors: list[str] = []
    indexed_files: list[str] = []

    for path in files:
        payload, err = load_single_document(path)
        if err:
            errors.append(f"{path.name}: {err}")
            continue
        indexed_files.append(path.name)
        text = payload["text"][:MAX_CHARS_PER_DOC]
        meta = payload["metadata"]
        for piece in chunk_text_simple(text):
            chunks.append(piece)
            metadatas.append(meta)

    if not chunks:
        return None, indexed_files, errors or ["No readable text found in documents."]

    engine = GroqTfidfChatEngine(config, chunks, metadatas)
    return engine, indexed_files, errors


def _is_rate_limit_error(exc: Exception) -> bool:
    name = type(exc).__name__
    if name in {"RateLimitError", "RateLimitReached", "InsufficientQuotaError"}:
        return True
    msg = str(exc).lower()
    return "rate limit" in msg or "429" in msg or "quota" in msg


def _build_groq_embed_model(config: dict[str, str]):
    """FastEmbed for cloud; HuggingFace fallback if FastEmbed fails on Streamlit."""
    from llama_index.embeddings.fastembed import FastEmbedEmbedding

    try:
        return FastEmbedEmbedding(model_name=config["fastembed_model"])
    except Exception as exc:
        logger.warning("FastEmbed unavailable (%s); using HuggingFace fallback.", exc)
        from llama_index.embeddings.huggingface import HuggingFaceEmbedding

        return HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")


def build_index_for_provider(documents: list, provider: str) -> Any:
    """Build vector index with provider-specific strategy."""
    from llama_index.core import VectorStoreIndex

    if provider == "openai":
        return build_index_resilient(documents, use_openai=True)

    # Groq / Ollama: index one document at a time (lower memory, more stable on cloud)
    index = VectorStoreIndex.from_documents([documents[0]])
    for doc in documents[1:]:
        index.insert(doc)
    return index


def build_index_resilient(documents: list, use_openai: bool) -> Any:
    """Build vector index; retries only needed for OpenAI rate limits."""
    from llama_index.core import VectorStoreIndex

    if not use_openai:
        if len(documents) == 1:
            return VectorStoreIndex.from_documents(documents)
        index = VectorStoreIndex.from_documents([documents[0]])
        for doc in documents[1:]:
            index.insert(doc)
        return index

    last_err: Exception | None = None
    for attempt in range(15):
        try:
            if len(documents) == 1:
                return VectorStoreIndex.from_documents(documents)
            index = VectorStoreIndex.from_documents([documents[0]])
            for doc in documents[1:]:
                time.sleep(5)
                index.insert(doc)
            return index
        except Exception as exc:
            if not _is_rate_limit_error(exc):
                raise
            last_err = exc
            wait = min(120, 15 * (attempt + 1))
            logger.warning("OpenAI rate limit — waiting %ds (attempt %d/15)", wait, attempt + 1)
            time.sleep(wait)

    raise last_err or RuntimeError("Indexing failed after repeated rate limits.")


@st.cache_resource(show_spinner=False)
def get_vector_index(model_key: str, corpus_fingerprint: str, config_snapshot: tuple):
    """Build or load a persisted vector index."""
    from llama_index.core import StorageContext, VectorStoreIndex, load_index_from_storage

    config = dict(config_snapshot)
    configure_llama_settings(config)
    provider = config["provider"]

    files = scan_data_source()
    if not files:
        return None, [], []

    documents, errors = load_all_documents(files)
    if not documents:
        return None, [], errors

    # Streamlit Cloud + Groq: always build in-memory (no disk cache — avoids ValueError)
    if provider == "groq":
        index = build_index_for_provider(documents, provider)
        return index, [d.metadata["filename"] for d in documents], errors

    INDEX_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_fp = f"{corpus_fingerprint}:{model_key}"
    if (
        FINGERPRINT_FILE.exists()
        and PERSIST_DIR.exists()
        and FINGERPRINT_FILE.read_text(encoding="utf-8").strip() == cache_fp
    ):
        try:
            storage_context = StorageContext.from_defaults(persist_dir=str(PERSIST_DIR))
            index = load_index_from_storage(storage_context)
            return index, [p.name for p in files], []
        except Exception:
            logger.warning("Could not load saved index; rebuilding.")
            shutil.rmtree(PERSIST_DIR, ignore_errors=True)

    index = build_index_for_provider(documents, provider)
    index.storage_context.persist(persist_dir=str(PERSIST_DIR))
    FINGERPRINT_FILE.write_text(cache_fp, encoding="utf-8")

    return index, [d.metadata["filename"] for d in documents], errors


def extract_answer_text(response: Any) -> str:
    """Normalize LlamaIndex chat response objects to plain text."""
    if response is None:
        raise ValueError("AI returned an empty response.")

    if hasattr(response, "response") and response.response:
        return str(response.response).strip()

    if hasattr(response, "message"):
        msg = response.message
        if hasattr(msg, "content") and msg.content:
            return str(msg.content).strip()

    text = str(response).strip()
    if not text or text == "None":
        raise ValueError("AI returned an empty response.")
    return text


def create_chat_engine(index, config: dict[str, str], history: list[dict[str, str]] | None = None):
    from llama_index.core import Settings
    from llama_index.core.chat_engine import CondensePlusContextChatEngine, ContextChatEngine
    from llama_index.core.llms import ChatMessage, MessageRole
    from llama_index.core.memory import ChatMemoryBuffer

    configure_llama_settings(config)
    memory = ChatMemoryBuffer.from_defaults(token_limit=3000)
    for msg in history or []:
        role = MessageRole.USER if msg["role"] == "user" else MessageRole.ASSISTANT
        memory.put(ChatMessage(role=role, content=msg["content"]))

    retriever = index.as_retriever(similarity_top_k=3 if config["provider"] == "groq" else 5)

    # Groq works more reliably with the simpler context chat engine
    if config["provider"] == "groq":
        return ContextChatEngine.from_defaults(
            retriever=retriever,
            llm=Settings.llm,
            memory=memory,
            system_prompt=SYSTEM_PROMPT,
            verbose=False,
        )

    return CondensePlusContextChatEngine.from_defaults(
        retriever=retriever,
        llm=Settings.llm,
        memory=memory,
        system_prompt=SYSTEM_PROMPT,
        verbose=False,
    )


def format_citation_html(response_text: str) -> str:
    """Wrap citation lines in a styled callout if present."""
    lines = response_text.strip().splitlines()
    citation_lines: list[str] = []
    body_lines: list[str] = []
    in_citations = True

    for line in lines:
        stripped = line.strip()
        if in_citations and ("🏛️" in stripped or "Organization:" in stripped):
            citation_lines.append(stripped)
        elif in_citations and stripped == "" and not citation_lines:
            continue
        else:
            in_citations = False
            body_lines.append(line)

    if citation_lines:
        citations_html = "<br>".join(citation_lines)
        body_html = "<br>".join(body_lines).replace("\n", "<br>")
        return (
            f'<div class="citation-box"><strong>📋 Official Source Citations</strong><br><br>'
            f"{citations_html}</div>{body_html}"
        )
    return response_text.replace("\n", "<br>")


def unique_source_records(metas: Any) -> list[dict[str, str]]:
    """Deduplicate source documents by filename while preserving citation metadata."""
    sources: list[dict[str, str]] = []
    seen: set[str] = set()
    for meta in metas:
        filename = (meta.get("filename") or "").strip()
        if not filename or filename in seen:
            continue
        seen.add(filename)
        sources.append(
            {
                "filename": filename,
                "organization": meta.get("organization", "Not explicitly mentioned in text"),
                "year": meta.get("year", "Not explicitly mentioned in text"),
                "reference_section": meta.get("reference_section", "Not explicitly mentioned in text"),
            }
        )
    return sources


def extract_sources_from_response(response: Any) -> list[dict[str, str]]:
    """Collect source filenames from LlamaIndex chat responses."""
    nodes = getattr(response, "source_nodes", None) or []
    metas: list[dict[str, str]] = []
    for item in nodes:
        node = getattr(item, "node", item)
        meta = getattr(node, "metadata", None) or {}
        if meta:
            metas.append(meta)
    return unique_source_records(metas)


def resolve_document_path(filename: str) -> Path | None:
    path = DATA_SOURCE_DIR / filename
    if path.is_file():
        return path
    return None


def mime_type_for_path(path: Path) -> str:
    suffix = path.suffix.lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }.get(suffix, "application/octet-stream")


@st.cache_data(show_spinner=False)
def read_document_bytes(path_str: str) -> bytes:
    return Path(path_str).read_bytes()


def render_pdf_preview(path: Path, key_prefix: str) -> None:
    """Inline PDF viewer — uses streamlit[pdf] when available, else browser iframe."""
    pdf_bytes = read_document_bytes(str(path))
    if len(pdf_bytes) > 12_000_000:
        st.caption("This document is large. Download it to view the full file locally.")
        return

    if hasattr(st, "pdf"):
        try:
            st.pdf(pdf_bytes, height=560)
            return
        except Exception as exc:
            logger.warning("st.pdf unavailable (%s); using iframe fallback.", exc)

    encoded = base64.b64encode(pdf_bytes).decode("ascii")
    st.components.v1.html(
        (
            f'<iframe src="data:application/pdf;base64,{encoded}" '
            f'width="100%" height="560" style="border:1px solid #D4DCE8;border-radius:8px;"></iframe>'
        ),
        height=580,
        scrolling=False,
    )


def render_source_documents(
    sources: list[dict[str, str]],
    key_prefix: str,
    *,
    expand_first: bool = False,
) -> None:
    """Show related source documents with optional inline PDF preview and download."""
    if not sources:
        return

    st.markdown('<p class="source-docs-header">📄 Related source documents</p>', unsafe_allow_html=True)
    st.caption("Expand a document below to preview the PDF in your browser or download the original file.")

    for index, source in enumerate(sources):
        filename = source.get("filename", "")
        path = resolve_document_path(filename)
        if path is None:
            st.caption(f"Source file not found on disk: {filename}")
            continue

        section = source.get("reference_section", "Not explicitly mentioned in text")
        expander_label = f"{filename}"
        if section and section != "Not explicitly mentioned in text":
            expander_label += f" — {section}"

        with st.expander(expander_label, expanded=expand_first and index == 0):
            st.markdown(
                (
                    f'<p class="source-doc-meta">'
                    f"🏛️ {source.get('organization', 'Government Authority')} · "
                    f"📅 {source.get('year', 'Not explicitly mentioned in text')}"
                    f"</p>"
                ),
                unsafe_allow_html=True,
            )

            file_bytes = read_document_bytes(str(path))
            st.download_button(
                label="⬇️ Download original document",
                data=file_bytes,
                file_name=filename,
                mime=mime_type_for_path(path),
                key=f"{key_prefix}_download_{index}_{filename}",
                use_container_width=True,
            )

            suffix = path.suffix.lower()
            if suffix == ".pdf":
                st.markdown("**Preview**")
                render_pdf_preview(path, key_prefix=f"{key_prefix}_pdf_{index}")
            elif suffix in {".docx", ".doc"}:
                st.info("Word documents can be downloaded and opened locally. PDF preview is shown when the source file is a PDF.")
            else:
                preview = path.read_text(encoding="utf-8", errors="replace")[:4000]
                st.text_area("Document excerpt", preview, height=220, disabled=True)


# ---------------------------------------------------------------------------
# Session state
# ---------------------------------------------------------------------------
def init_session_state() -> None:
    defaults = {
        "messages": [],
        "chat_engine": None,
        "admin_unlocked": False,
        "corpus_fingerprint": "",
        "indexed_files": [],
        "ingest_errors": [],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def reset_conversation() -> None:
    st.session_state.messages = []
    st.session_state.chat_engine = None


# ---------------------------------------------------------------------------
# UI components
# ---------------------------------------------------------------------------
def render_header() -> None:
    st.markdown(
        f"""
        <div class="gov-header">
            <a class="gov-site-fab" href="{DEVELOPER_SITE_URL}" target="_blank" rel="noopener noreferrer"
               title="Visit tharshan.lk">
                <span class="gov-site-fab-icon">🌐</span>
                <span>tharshan.lk</span>
            </a>
            <h1>🏛️ {APP_NAME}</h1>
            <p>{APP_TAGLINE}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def inject_cloud_toolbar_guard() -> None:
    """Custom Share modal; block Streamlit/GitHub/profile links; hide empty menu."""
    ui_path = BASE_DIR / "scripts" / "gov-ui.js"
    ui_script = ui_path.read_text(encoding="utf-8").replace("__APP_TITLE__", APP_NAME)
    st.components.v1.html(
        f"<script>{ui_script}</script>",
        height=0,
        width=0,
        scrolling=False,
    )


def render_footer() -> None:
    year = datetime.now().year
    st.markdown(
        f"""
        <div class="gov-footer">
            <div class="gov-footer-inner">
                <span class="gov-footer-label">Developed by</span>
                <span class="gov-footer-name">S. THARSHAN</span>
                <p class="gov-footer-tagline">"Building Intelligent Solutions for Tomorrow"</p>
                <a class="gov-footer-email" href="mailto:tharsh.ai.dev@gmail.com">
                    ✉️ tharsh.ai.dev@gmail.com
                </a>
                <div class="gov-footer-divider"></div>
                <p class="gov-footer-copy">© {year} {APP_NAME} · tharshan.lk</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_status_badge(label: str, ok: bool, detail: str = "") -> None:
    css_class = "status-ok" if ok else "status-error"
    icon = "✅" if ok else "⚠️"
    detail_html = f" — {detail}" if detail else ""
    st.markdown(
        f'<span class="status-badge {css_class}">{icon} {label}{detail_html}</span>',
        unsafe_allow_html=True,
    )


def config_snapshot(config: dict[str, str]) -> tuple[tuple[str, str], ...]:
    return tuple(sorted(config.items()))


def render_sidebar(config: dict[str, str], files: list[Path]) -> None:
    with st.sidebar:
        st.markdown('<p class="sidebar-title">System Status</p>', unsafe_allow_html=True)

        if config["provider"] == "openai":
            render_status_badge("OpenAI API Key", bool(config["openai_api_key"]))
            st.caption(f"Model: {config['openai_model']}")
        elif config["provider"] == "groq":
            render_status_badge("Groq API Key", bool(config["groq_api_key"]), "Free cloud AI")
            st.caption(f"Model: {config['groq_model']}")
        else:
            ollama_ok = ollama_is_running(config["ollama_base_url"])
            render_status_badge("Ollama Running", ollama_ok, "Free local AI")
            st.caption(f"Model: {config['ollama_model']}")

        render_status_badge("Knowledge Base Ready", is_ai_ready(config) and len(files) > 0, "Indexed")
        st.markdown(
            f'<span class="status-badge status-ok">📄 Active Documents: {len(files)}</span>',
            unsafe_allow_html=True,
        )
        if st.session_state.indexed_files:
            st.caption(f"Indexed: {len(st.session_state.indexed_files)} file(s)")

        if documents_changed(files) and st.session_state.chat_engine is not None:
            st.markdown(
                '<span class="status-badge status-warn">🔄 New/updated documents — re-indexing…</span>',
                unsafe_allow_html=True,
            )
        else:
            st.caption("Add PDFs to `data_source/` — indexed automatically on refresh.")

        if st.button("🔄 Refresh documents", use_container_width=True, help="Re-scan data_source/ after adding files"):
            st.rerun()

        if st.session_state.ingest_errors:
            st.markdown(
                f'<span class="status-badge status-warn">⚠️ {len(st.session_state.ingest_errors)} file(s) skipped</span>',
                unsafe_allow_html=True,
            )

        st.divider()
        st.markdown('<p class="sidebar-title">Session</p>', unsafe_allow_html=True)
        st.caption(f"Messages: {len(st.session_state.messages)}")
        if st.button("🔄 Clear Conversation", use_container_width=True):
            reset_conversation()
            st.rerun()

        st.divider()
        st.markdown('<p class="sidebar-title">Administrator</p>', unsafe_allow_html=True)

        if st.session_state.admin_unlocked:
            st.success("Admin access granted")
            with st.expander("📁 Document Management Guide", expanded=True):
                st.markdown(
                    f"""
                    **Adding new documents**

                    1. Place PDF, DOCX, or TXT files in:
                       `{DATA_SOURCE_DIR.name}/`
                    2. Use descriptive filenames, e.g.  
                       `Finance_Ministry_Circular_2023.pdf`
                    3. **Refresh the page** or click **Refresh documents** in the sidebar.
                       The knowledge base re-indexes automatically — no manual rebuild needed.
                    4. Use **Rebuild Knowledge Base** only if indexing seems stuck.

                    **Supported formats:** PDF, DOCX, TXT, MD  
                    **Indexed files:** {len(st.session_state.indexed_files)}

                    **Deployment note:** Upload files to `data_source/` in your
                    repository, then reboot the app or click **Refresh documents**.
                    """
                )
                if st.session_state.ingest_errors:
                    st.warning("Files with errors:")
                    for err in st.session_state.ingest_errors:
                        st.caption(f"• {err}")

            if st.button("🔨 Rebuild Knowledge Base", use_container_width=True):
                get_vector_index.clear()
                get_groq_tfidf_engine.clear()
                reset_conversation()
                st.session_state.corpus_fingerprint = ""
                if INDEX_CACHE_DIR.exists():
                    shutil.rmtree(INDEX_CACHE_DIR, ignore_errors=True)
                st.rerun()

            if st.button("🔒 Lock Admin Panel", use_container_width=True):
                st.session_state.admin_unlocked = False
                st.rerun()
        else:
            pwd = st.text_input("Admin password", type="password", key="admin_pwd_input")
            if st.button("Unlock Admin Panel", use_container_width=True):
                if pwd == config["admin_password"]:
                    st.session_state.admin_unlocked = True
                    st.rerun()
                elif pwd:
                    st.error("Incorrect password.")

        st.divider()
        st.caption(f"© {datetime.now().year} · tharshan.lk · {config['provider'].upper()}")


def _persist_index_fingerprint(cache_fp: str) -> None:
    INDEX_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    FINGERPRINT_FILE.write_text(cache_fp, encoding="utf-8")


def _index_cache_matches(fingerprint: str, model_key: str) -> bool:
    cache_fp = f"{fingerprint}:{model_key}"
    return (
        FINGERPRINT_FILE.exists()
        and FINGERPRINT_FILE.read_text(encoding="utf-8").strip() == cache_fp
    )


def restore_chat_engine_from_cache(config: dict[str, str], fingerprint: str) -> bool:
    """Reuse cached knowledge base without re-running indexing UI."""
    mkey = model_cache_key(config)
    has_disk_cache = _index_cache_matches(fingerprint, mkey)
    session_matches = st.session_state.corpus_fingerprint == fingerprint

    if st.session_state.chat_engine is not None and session_matches:
        return True

    if not session_matches and not has_disk_cache:
        return False

    if config["provider"] == "groq":
        engine, indexed_files, errors = get_groq_tfidf_engine(
            fingerprint,
            config_snapshot(config),
        )
        if engine is None:
            return False
        st.session_state.chat_engine = engine
        st.session_state.indexed_files = indexed_files
        st.session_state.ingest_errors = errors
        st.session_state.corpus_fingerprint = fingerprint
        return True

    if not (has_disk_cache and PERSIST_DIR.exists()):
        return False

    index, indexed_files, errors = get_vector_index(
        mkey,
        fingerprint,
        config_snapshot(config),
    )
    if index is None:
        return False
    st.session_state.chat_engine = create_chat_engine(
        index,
        config,
        history=st.session_state.messages,
    )
    st.session_state.indexed_files = indexed_files
    st.session_state.ingest_errors = errors
    st.session_state.corpus_fingerprint = fingerprint
    return True


def initialize_rag(config: dict[str, str], files: list[Path]) -> bool:
    if not is_ai_ready(config):
        return False
    if not files:
        return False

    fingerprint = compute_corpus_fingerprint(files)
    if restore_chat_engine_from_cache(config, fingerprint):
        return True

    if st.session_state.corpus_fingerprint and st.session_state.corpus_fingerprint != fingerprint:
        get_vector_index.clear()
        get_groq_tfidf_engine.clear()
        st.session_state.chat_engine = None

    try:
        mkey = model_cache_key(config)
        cached = _index_cache_matches(fingerprint, mkey)

        provider_label = {
            "openai": "OpenAI",
            "groq": "Groq (free)",
            "ollama": "Ollama (local)",
        }.get(config["provider"], config["provider"])

        with st.status(
            "Loading knowledge base from cache…" if cached else "Indexing government documents…",
            expanded=not cached,
        ) as status:
            if not cached:
                st.write("📖 **Step 1/2** — Reading your PDF files…")
                for path in files:
                    st.caption(f"• {path.name}")
                timing = (
                    "**First time only:** indexing usually takes **1–3 minutes**. "
                    "Scanned PDFs use Groq Vision OCR and may take a little longer."
                    if config["provider"] == "groq"
                    else "**First time only:** large manuals can take **10–30 minutes** on Ollama. Keep this tab open."
                )
                st.write(f"🔄 **Step 2/2** — Indexing with **{provider_label}**.\n\n{timing}")
            else:
                st.write("✅ Using saved index — no re-indexing needed.")

            if config["provider"] == "groq":
                engine, indexed_files, errors = get_groq_tfidf_engine(
                    fingerprint,
                    config_snapshot(config),
                )
                st.session_state.ingest_errors = errors
                st.session_state.indexed_files = indexed_files
                st.session_state.corpus_fingerprint = fingerprint
                if engine is None:
                    st.session_state.chat_engine = None
                    status.update(label="Indexing failed", state="error")
                    detail = "; ".join(errors[:3]) if errors else "No readable text in documents."
                    st.error(f"Unable to build the knowledge base. {detail}")
                    return False
                st.session_state.chat_engine = engine
                _persist_index_fingerprint(f"{fingerprint}:{mkey}")
                status.update(label="Indexing complete — ready to chat", state="complete")
                return True

            index, indexed_files, errors = get_vector_index(
                mkey,
                fingerprint,
                config_snapshot(config),
            )

            if cached:
                status.update(label="Knowledge base loaded from cache", state="complete")
            else:
                status.update(label="Indexing complete — ready to chat", state="complete")
        st.session_state.ingest_errors = errors
        st.session_state.indexed_files = indexed_files
        st.session_state.corpus_fingerprint = fingerprint

        if index is None:
            st.session_state.chat_engine = None
            return False

        st.session_state.chat_engine = create_chat_engine(
            index,
            config,
            history=st.session_state.messages,
        )
        return True
    except Exception as exc:
        logger.exception("RAG initialization failed")
        st.session_state.chat_engine = None
        if _is_rate_limit_error(exc):
            st.error("OpenAI rate limit or quota reached while indexing documents.")
            st.markdown(
                """
                **Switch to free local AI instead:**

                1. Install [Ollama](https://ollama.com/download)
                2. In terminal run:
                   ```
                   ollama pull llama3.2
                   ollama pull nomic-embed-text
                   ```
                3. In `.streamlit/secrets.toml` set: `AI_PROVIDER = "ollama"`
                4. Restart the app

                Or add billing at [platform.openai.com/settings/billing](https://platform.openai.com/settings/billing).
                """
            )
        else:
            st.error(f"Unable to initialize the knowledge base. ({type(exc).__name__})")
            st.caption(str(exc)[:500])
            if st.session_state.admin_unlocked:
                st.code(str(exc))
            st.markdown(
                """
                **Try these fixes:**
                1. Streamlit Cloud → **Manage app** → **Reboot app**
                2. Sidebar (Admin) → **Rebuild Knowledge Base**
                3. Confirm Secrets: `AI_PROVIDER = "groq"` and valid `GROQ_API_KEY`
                4. Push latest code to your repository, then reboot again
                """
            )
        return False


def render_local_setup_help(config: dict[str, str]) -> None:
    """Show localhost setup steps when the AI provider is not configured."""
    if config["provider"] == "groq" and not config["groq_api_key"]:
        st.error("Groq API key is missing on localhost.")
        st.markdown(
            """
            **Local setup (one time):**

            1. Copy `.streamlit/secrets.toml.example` → `.streamlit/secrets.toml`
            2. Paste your free Groq key from [console.groq.com](https://console.groq.com)
            3. Restart Streamlit (`Ctrl+C`, then `streamlit run app.py` again)

            Example `.streamlit/secrets.toml`:
            ```
            AI_PROVIDER = "groq"
            GROQ_API_KEY = "gsk_your_real_key_here"
            GROQ_MODEL = "llama-3.1-8b-instant"
            ADMIN_PASSWORD = "admin123"
            ```
            """
        )
    elif config["provider"] == "openai" and not config["openai_api_key"]:
        st.error("OpenAI API key is missing. Add `OPENAI_API_KEY` to `.streamlit/secrets.toml` or `.env`.")
    elif config["provider"] == "ollama" and not ollama_is_running(config["ollama_base_url"]):
        st.error("Ollama is not running. Start Ollama, pull models, then refresh this page.")


def chat_is_ready(config: dict[str, str]) -> bool:
    return is_ai_ready(config) and st.session_state.chat_engine is not None


def process_user_question(config: dict[str, str], prompt: str) -> None:
    """Run one officer query and append assistant/user messages to session state."""
    prompt = prompt.strip()
    if not prompt:
        return

    if not is_ai_ready(config):
        st.warning("AI engine is not ready. Complete local setup steps above.")
        return

    if st.session_state.chat_engine is None:
        st.warning(
            "Knowledge base is not ready yet. Wait for indexing to finish, "
            "or use **Rebuild Knowledge Base** in the sidebar."
        )
        return

    st.session_state.messages.append({"role": "user", "content": prompt})

    sources: list[dict[str, str]] = []
    with st.spinner("Searching official documents…"):
        try:
            if config["provider"] == "groq" and hasattr(st.session_state.chat_engine, "chat"):
                answer, sources = st.session_state.chat_engine.chat(
                    prompt,
                    history=st.session_state.messages[:-1],
                )
            else:
                response = st.session_state.chat_engine.chat(prompt)
                answer = extract_answer_text(response)
                sources = extract_sources_from_response(response)
        except Exception as exc:
            logger.exception("Chat request failed")
            answer = (
                "I apologize — a temporary error occurred while processing your request. "
                "Please try again in a moment. If the issue persists, contact your system administrator."
            )
            if st.session_state.admin_unlocked:
                st.caption(f"Error detail (admin): {type(exc).__name__}: {exc}")
            else:
                st.caption(f"Error detail (admin): {type(exc).__name__}")

    st.session_state.messages.append(
        {"role": "assistant", "content": answer, "sources": sources}
    )
    st.rerun()


def render_chat(config: dict[str, str]) -> None:
    ready = chat_is_ready(config)

    st.markdown('<div class="query-panel">', unsafe_allow_html=True)
    st.markdown("#### 💬 Ask your question")
    if ready:
        indexed_list = ", ".join(st.session_state.indexed_files) or "none"
        st.caption(
            f"**{len(st.session_state.indexed_files)}** of **{len(scan_data_source())}** document(s) indexed: {indexed_list}. "
            "After you search, expand **Related source documents** to preview or download PDFs."
        )
        if st.session_state.ingest_errors:
            for err in st.session_state.ingest_errors:
                st.warning(f"⚠️ Not indexed: {err}")
    else:
        st.caption("Complete setup and indexing before searching documents.")

    with st.form("officer_query_form", clear_on_submit=True):
        st.markdown('<div class="query-row">', unsafe_allow_html=True)
        col_input, col_button = st.columns([5, 1])
        with col_input:
            prompt = st.text_input(
                "Officer question",
                placeholder="Ask about circulars, guidelines, or regulations…",
                disabled=not ready,
                label_visibility="collapsed",
            )
        with col_button:
            submitted = st.form_submit_button(
                "Ask AI",
                type="primary",
                disabled=not ready,
                use_container_width=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)
        st.caption("Press **Enter** or click **Ask AI** to search.")
    st.markdown("</div>", unsafe_allow_html=True)

    if submitted and prompt.strip():
        process_user_question(config, prompt)

    for msg_index, message in enumerate(st.session_state.messages):
        with st.chat_message(message["role"]):
            if message["role"] == "assistant":
                st.markdown(format_citation_html(message["content"]), unsafe_allow_html=True)
                render_source_documents(
                    message.get("sources", []),
                    key_prefix=f"history_{msg_index}",
                    expand_first=msg_index == len(st.session_state.messages) - 1,
                )
            else:
                st.markdown(message["content"])


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    st.set_page_config(
        page_title=APP_NAME,
        page_icon="🏛️",
        layout="wide",
        initial_sidebar_state="expanded",
        menu_items={
            "Get help": None,
            "Report a bug": None,
            "About": None,
        },
    )
    st.markdown(GOVERNMENT_CSS, unsafe_allow_html=True)

    init_session_state()
    config = get_config()
    files = scan_data_source()

    render_sidebar(config, files)
    render_header()

    try:
        if not is_ai_ready(config):
            render_local_setup_help(config)
            render_chat(config)
            return

        if not files:
            st.warning(
                f"No documents found in `{DATA_SOURCE_DIR.name}/`. "
                "Administrators can add PDF or Word files and rebuild the knowledge base from the sidebar."
            )
            st.markdown(
                """
                **Example queries once documents are loaded:**
                - *What are the leave entitlements under the latest public service circular?*
                - *Summarize procurement guidelines for values above the threshold limit.*
                - *Which section covers disciplinary procedures?*
                """
            )
            return

        if not initialize_rag(config, files):
            st.warning(
                "Knowledge base could not be loaded. Check the error above, then use "
                "**Rebuild Knowledge Base** in the sidebar after fixing the issue."
            )

        render_chat(config)
    finally:
        render_footer()
        inject_cloud_toolbar_guard()


if __name__ == "__main__":
    main()
